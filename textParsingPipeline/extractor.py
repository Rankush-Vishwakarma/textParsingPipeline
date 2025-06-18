import json
from pathlib import Path
from typing import Union, BinaryIO
import pandas as pd
import io
import logging
from bs4 import BeautifulSoup
import fitz
from docx import Document

def extract_text_from_html(content):
    """Extract text from HTML content using BeautifulSoup."""
    soup = BeautifulSoup(content.decode('utf-8', errors='ignore'), 'html.parser')
    for script_or_style in soup(['script', 'style']):
        script_or_style.decompose()
    text = soup.get_text(separator=' ', strip=True)
    return text


def extract_text_from_pdf(content):
    """Extract text from PDF content using PyMuPDF."""
    with io.BytesIO(content) as f:
        doc = fitz.open(stream=f, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        return text


def extract_text_from_docx(content):
    """Extract text from DOCX content using python-docx."""
    with io.BytesIO(content) as f:
        doc = Document(f)
        text = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                text.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(text)
        return text

def extract_text_from_xlsx(content):
    try:
        excel = pd.read_excel(io.BytesIO(content), sheet_name=None)
        if not excel:
            return "No sheets found in the Excel file."
        return "\n".join([df.to_string(index=False) for df in excel.values()])
    except Exception as e:
        return f"Failed to extract XLSX content: {str(e)}"


def extract_text(content, filename):
    ext = filename.split('.')[-1].lower()
    if ext == 'pdf': return extract_text_from_pdf(content)
    if ext in ['docx', 'doc']: return extract_text_from_docx(content)
    if ext == 'html': return extract_text_from_html(content)
    if ext == 'xlsx': return extract_text_from_xlsx(content)
    raise ValueError("Unsupported file type")